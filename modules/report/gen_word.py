"""Context docxtpl + sinh báo cáo Word cho luồng "Tạo báo cáo Word".

API chính:
* ``mba(doc, ...)``, ``device(doc, ...)`` — dựng context cho từng template.
* ``merge_rendered_docx`` / ``merge_mba_device_docx`` — ghép nhiều file đã render.
* ``mba_kwargs_from_inps`` / ``mba_kwargs_from_folder`` / ``device_kwargs_from_folder`` —
  tự chọn ảnh từ thư mục thiết bị (**bắt buộc** ``a.png`` + PS-SDxxxx.BMP). Với MBA,
  ``mba_kwargs_from_folder`` tự tìm ``INPSxxxx.KEW`` (cùng quy ước ``find_file`` như phân tích KEW
  và như luồng Excel MBA) rồi gọi ``mba_kwargs_from_inps``; thiếu INPSxxxx.KEW thì bảng dùng ``"—"``.
  **Nhận xét văn bản** (``remarks_mba`` / ``remarks_device``): tự sinh theo dữ liệu từ file INPSxxxx.KEW; có thể ghép thêm ghi chú
  cột Excel (P, PF, …); nếu ô Excel chứa đoạn bắt đầu bằng ``Nhận xét:`` thì dùng nguyên văn thủ công.
* ``build_field_word_report`` — quét một thư mục ``Project_Output/`` rồi xuất
  1 file Word duy nhất gồm nhiều MBA / device.
* ``build_word_report_from_zip`` — entry-point cho API: nhận ZIP đã tổ chức
  (output của "Xử lý file sơ bộ"), tự dò metadata trong Excel kèm (nếu có),
  trả về đường dẫn báo cáo Word (bao gồm cả MBA, Chương 4 và Chương 5).
* ``build_chapter4_from_zip`` / ``build_chapter5_from_zip`` — Tách riêng Chương 4
  (chỉ thiết bị type=4) và Chương 5 (MBA + thiết bị còn lại).
* ``generate_table6_from_zip`` — tạo Bảng 6.3 tổng hợp.
"""

from __future__ import annotations

import io
import re
import unicodedata
import zipfile
from pathlib import Path
from shutil import copy2
from tempfile import TemporaryDirectory
from typing import Iterable, Literal, Mapping, Sequence

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Mm
from docxcompose.composer import Composer
from docxtpl import DocxTemplate, InlineImage

__all__ = [
    "mba",
    "device",
    "merge_rendered_docx",
    "merge_mba_device_docx",
    "mba_kwargs_from_inps",
    "mba_kwargs_from_folder",
    "device_kwargs_from_folder",
    "build_field_word_report",
    "build_word_report_from_zip",
    "build_chapter4_from_zip",
    "build_chapter5_from_zip",
    "DEFAULT_MBA_TEMPLATE",
    "DEFAULT_DEVICE_TEMPLATE",
    "DEFAULT_DEVICE4_TEMPLATE",
    "DEFAULT_TABLE6_TEMPLATE",
    "generate_table6_docx",
    "generate_table6_from_zip",
]

_REPO_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_MBA_TEMPLATE = _REPO_ROOT / "static" / "word-template" / "mba.docx"
DEFAULT_DEVICE_TEMPLATE = _REPO_ROOT / "static" / "word-template" / "device.docx"
DEFAULT_DEVICE4_TEMPLATE = _REPO_ROOT / "static" / "word-template" / "device4.docx"
DEFAULT_TOTALMBA_TEMPLATE = _REPO_ROOT / "static" / "word-template" / "totalmba.docx"
DEFAULT_TABLE6_TEMPLATE = _REPO_ROOT / "static" / "word-template" / "table6.docx"

WIDTH_LARGE, WIDTH_SMALL, HEIGHT_MBA = Mm(109.6), Mm(53.8), Mm(40.5)
WIDTH_A, WIDTH_SUB, HEIGHT_A, HEIGHT_SUB = Mm(166.3), Mm(54.3), Mm(60.0), Mm(41.3)

SectionKind = Literal["mba", "device", "device4", "total_mba"]

_DEVICE_KIND_LABELS = frozenset(
    {"device", "thiết bị", "thiet bi", "tủ", "tu", "khac", "khác"}
)
# Các cột dữ liệu ghi chú mở rộng từ hiện trường.
_EXCEL_METRIC_REMARKS: tuple[tuple[str, str], ...] = (
    ("p", "P"),
    ("pf", "PF"),
    ("i1", "I1"),
    ("i2", "I2"),
    ("i3", "I3"),
    ("di", "ΔI"),
    ("thd", "THD"),
    ("tdd", "TDD"),
)
# Các cột đặc tính hiện trường dùng để sinh nhận xét (thay thế INPS).
_FIELD_PARAM_COLS: tuple[str, ...] = (
    "current_char", "u_min", "u_max", "i_max",
    "delta_u", "delta_i", "p", "cos_phi",
    "thd", "tdd", "pdm",
)

# Ký tự sánh an toàn trong XML / docxtpl (thoát trước khi nạp vào Word template).
# Các chuỗi nhận xét dùng hằng này để biểu diễn "nhỏ hơn" / "lớn hơn" mà không bị
# trình phân tích XML của python-docx / docxtpl hiểu nhầm là thẻ XML.
_LT = "<"
_GT = ">"
_AMP = "&"
_V_DEV_LIMIT_PCT = 5.0
_PF_LIMIT = 0.9
_THDV_LIMIT_PCT = 8.0
_TDD_LIMIT_PCT = 12.0
# Điện áp danh định cho cột đánh giá lệch % U12: -5% ≤ δ ≤ +5% (so với 400 V), không lấy từ Excel/metadata.
_MBA_NOMINAL_VOLTAGE_V = 400.0
# Biến động dòng (INPS min/max/avg) — cùng ý tưởng với analyse_kew.CONFIG defaults.
_I_SPREAD_STABLE_PCT = 15.0
_I_SPREAD_MODERATE_PCT = 50.0
# Ảnh theo dải Excel / organize_field_zip: PS-SD641.BMP …
_BMP_RE = re.compile(r"^PS-SD(\d+)\.BMP$", re.IGNORECASE)

# ════════════════════════════════════════════════════════════════════
#                     Context builders (template)
# ════════════════════════════════════════════════════════════════════

def _inline(doc: DocxTemplate, path: str, height, width) -> InlineImage:
    """Tạo đối tượng InlineImage để chèn ảnh vào template Word.

    Args:
        doc: Đối tượng DocxTemplate đang xử lý.
        path: Đường dẫn vật lý đến file ảnh.
        height: Chiều cao ảnh (đơn vị Mm hoặc tương đương).
        width: Chiều rộng ảnh (đơn vị Mm hoặc tương đương).

    Returns:
        InlineImage: Đối tượng ảnh sẵn sàng để render vào template.
    """
    return InlineImage(doc, path, height=height, width=width)

def mba(
    doc: DocxTemplate,
    *,
    name: str,
    imga: str,
    img1: str,
    img2: str,
    img4: str,
    img6: str,
    cap_fig_mba: str,
    remarks_mba: str,
    cap_tab_mba: str,
    # Điện áp
    u12max: str,
    u12min: str,
    u12avg: str,
    u12eval: str,
    u23max: str,
    u23min: str,
    u23avg: str,
    u31max: str,
    u31min: str,
    u31avg: str,
    # Dòng điện
    i1max: str,
    i1min: str,
    i1avg: str,
    i2max: str,
    i2min: str,
    i2avg: str,
    i3max: str,
    i3min: str,
    i3avg: str,
    # Độ lệch pha điện áp / dòng & hệ số công suất
    dumax: str,
    dumin: str,
    duavg: str,
    dueval: str,
    dimax: str,
    dimin: str,
    diavg: str,
    pfmax: str,
    pfmin: str,
    pfavg: str,
    pfeval: str,
    # Công suất P, Q, S
    pmax: str,
    pmin: str,
    pavg: str,
    qmax: str,
    qmin: str,
    qavg: str,
    smax: str,
    smin: str,
    savg: str,
    # Sóng hài THD
    thd1max: str,
    thd1min: str,
    thd1avg: str,
    thdeval: str,
    thd2max: str,
    thd2min: str,
    thd2avg: str,
    thd3max: str,
    thd3min: str,
    thd3avg: str,
    # Sóng hài TDD
    tdd1max: str,
    tdd1min: str,
    tdd1avg: str,
    tddeval: str,
    tdd2max: str,
    tdd2min: str,
    tdd2avg: str,
    tdd3max: str,
    tdd3min: str,
    tdd3avg: str,
    # Danh sách thiết bị (tuỳ chọn)
    ds_mba: list[dict] | None = None,
    **kwargs,
) -> dict:
    """Xây dựng context dictionary đầy đủ cho template báo cáo Máy biến áp (MBA).

    Args:
        doc: Đối tượng DocxTemplate dùng để khởi tạo InlineImage.
        name: Tên hiển thị của máy biến áp.
        imga, img1, img2, img4, img6: Đường dẫn các file ảnh tương ứng (Overview, I, THD...).
        cap_fig_mba: Chú thích cho phần hình ảnh.
        remarks_mba: Nội dung nhận xét văn bản.
        cap_tab_mba: Chú thích cho bảng thông số.
        u12max, u12min, u12avg, u12eval: Các thông số điện áp dây pha 1-2.
        u23max, u23min, u23avg: Các thông số điện áp dây pha 2-3.
        u31max, u31min, u31avg: Các thông số điện áp dây pha 3-1.
        i1max...i3avg: Các thông số dòng điện 3 pha.
        dumax...dueval: Độ lệch pha điện áp và đánh giá.
        dimax...diavg: Độ lệch pha dòng điện.
        pfmax...pfeval: Hệ số công suất và đánh giá.
        pmax...savg: Công suất tác dụng (P), phản kháng (Q), biểu kiến (S).
        thd1max...thdeval: Sóng hài THD điện áp và đánh giá.
        tdd1max...tddeval: Sóng hài TDD dòng điện và đánh giá.
        ds_mba: Danh sách tóm tắt thiết bị cho bảng tổng hợp (nếu có).
        **kwargs: Các tham số mở rộng khác.

    Returns:
        dict: Tập hợp các cặp key-value để nạp vào template docxtpl.
    """
    return {
        "mba_name": name,
        "imga": _inline(doc, imga, HEIGHT_MBA, WIDTH_LARGE),
        "img1": _inline(doc, img1, HEIGHT_MBA, WIDTH_SMALL),
        "img2": _inline(doc, img2, HEIGHT_MBA, WIDTH_SMALL),
        "img4": _inline(doc, img4, HEIGHT_MBA, WIDTH_SMALL),
        "img6": _inline(doc, img6, HEIGHT_MBA, WIDTH_SMALL),
        "cap_fig_mba": cap_fig_mba,
        "remarks_mba": remarks_mba,
        "cap_tab_mba": cap_tab_mba,
        "u12max": u12max, "u12min": u12min, "u12avg": u12avg, "u12eval": u12eval,
        "u23max": u23max, "u23min": u23min, "u23avg": u23avg,
        "u31max": u31max, "u31min": u31min, "u31avg": u31avg,
        "i1max": i1max, "i1min": i1min, "i1avg": i1avg,
        "i2max": i2max, "i2min": i2min, "i2avg": i2avg,
        "i3max": i3max, "i3min": i3min, "i3avg": i3avg,
        "dumax": dumax, "dumin": dumin, "duavg": duavg, "dueval": dueval,
        "dimax": dimax, "dimin": dimin, "diavg": diavg,
        "pfmax": pfmax, "pfmin": pfmin, "pfavg": pfavg, "pfeval": pfeval,
        "pmax": pmax, "pmin": pmin, "pavg": pavg,
        "qmax": qmax, "qmin": qmin, "qavg": qavg,
        "smax": smax, "smin": smin, "savg": savg,
        "thd1max": thd1max, "thd1min": thd1min, "thd1avg": thd1avg, "thdeval": thdeval,
        "thd2max": thd2max, "thd2min": thd2min, "thd2avg": thd2avg,
        "thd3max": thd3max, "thd3min": thd3min, "thd3avg": thd3avg,
        "tdd1max": tdd1max, "tdd1min": tdd1min, "tdd1avg": tdd1avg, "tddeval": tddeval,
        "tdd2max": tdd2max, "tdd2min": tdd2min, "tdd2avg": tdd2avg,
        "tdd3max": tdd3max, "tdd3min": tdd3min, "tdd3avg": tdd3avg,
        "ds_mba": ds_mba or [],
        **kwargs,
    }

def device(
    doc: DocxTemplate,
    *,
    name: str,
    imga: str,
    img1: str,
    img2: str,
    img3: str,
    img4: str,
    img5: str,
    img6: str,
    cap_device: str,
    remarks_device: str,
) -> dict:
    """Xây dựng context dictionary cho template báo cáo thiết bị phụ tải (Chương 4/5).

    Args:
        doc: Đối tượng DocxTemplate dùng để khởi tạo InlineImage.
        name: Tên thiết bị đo kiểm.
        imga: Đường dẫn ảnh tổng quan (thường là a.png).
        img1, img2, img3, img4, img5, img6: Đường dẫn tới 6 ảnh thông số (BMP).
        cap_device: Chú thích bảng/hình ảnh.
        remarks_device: Nội dung nhận xét văn bản cho thiết bị.

    Returns:
        dict: Context dữ liệu để render template.
    """
    return {
        "device_name": name,
        "imga": _inline(doc, imga, HEIGHT_A, WIDTH_A),
        "img1": _inline(doc, img1, HEIGHT_SUB, WIDTH_SUB),
        "img2": _inline(doc, img2, HEIGHT_SUB, WIDTH_SUB),
        "img3": _inline(doc, img3, HEIGHT_SUB, WIDTH_SUB),
        "img4": _inline(doc, img4, HEIGHT_SUB, WIDTH_SUB),
        "img5": _inline(doc, img5, HEIGHT_SUB, WIDTH_SUB),
        "img6": _inline(doc, img6, HEIGHT_SUB, WIDTH_SUB),
        "cap1": cap_device,
        "remarks": remarks_device,
    }

def total_mba(doc: DocxTemplate, *, ds_mba: list[dict], **kwargs) -> dict:
    """Dựng context cho template bảng tổng hợp danh sách máy biến áp.

    Args:
        doc: Đối tượng DocxTemplate.
        ds_mba: Danh sách các dict chứa thông tin rút gọn (stt, tên, công suất, cosphi).
        **kwargs: Các tham số mở rộng.

    Returns:
        dict: Context chứa khóa 'ds_mba'.
    """
    return {"ds_mba": ds_mba}

def _format_name_mid(name: str) -> str:
    """Định dạng tên thiết bị khi đưa vào giữa câu văn.

    Quy tắc:
    - Nếu từ đầu tiên viết hoa toàn bộ (VD: MBA, MSB, T1...), giữ nguyên hoa.
    - Ngược lại, chuyển chữ cái đầu tiên sang chữ thường (VD: "Máy bơm" -> "máy bơm").

    Args:
        name: Tên gốc của thiết bị.

    Returns:
        str: Tên đã định dạng lại phù hợp ngữ cảnh giữa câu.
    """
    if not name:
        return ""
    w = name.split()
    first_upper = w[0].isupper() if w else False
    if first_upper:
        return name
    return name[0].lower() + name[1:]

# ════════════════════════════════════════════════════════════════════
#                            Format helpers
# ════════════════════════════════════════════════════════════════════

def _f(v, d: int = 1) -> str:
    """Định dạng giá trị số sang chuỗi theo kiểu Việt Nam (dấu phẩy thập phân).

    Args:
        v: Giá trị cần định dạng (int, float, str hoặc None).
        d: Số chữ số thập phân sau dấu phẩy. Mặc định là 1.

    Returns:
        str: Chuỗi số đã định dạng (VD: "220,5"). Trả về "—" nếu v là None hoặc không hợp lệ.
    """
    if v is None:
        return "—"
    try:
        x = float(v)
    except (TypeError, ValueError):
        return "—"
    if x != x:  # NaN
        return "—"
    return f"{x:.{d}f}".replace(".", ",")

def _tri(d: Mapping, dec: int = 1) -> tuple[str, str, str]:
    """Chuyển đổi bộ ba thống kê (min, max, avg) sang chuỗi đã định dạng.

    Args:
        d: Mapping chứa các khóa 'min', 'max', 'avg'.
        dec: Số chữ số thập phân cho hàm định dạng _f. Mặc định là 1.

    Returns:
        tuple[str, str, str]: Bộ ba chuỗi (max, min, avg) đã định dạng dấu phẩy.
    """
    return _f(d.get("max"), dec), _f(d.get("min"), dec), _f(d.get("avg"), dec)

# ════════════════════════════════════════════════════════════════════
#                       INPS aggregation helpers
# ════════════════════════════════════════════════════════════════════

def _parse_inps(inps_path: str | Path) -> dict[str, dict]:
    """Đọc và tóm tắt dữ liệu từ file INPSxxxx.KEW.

    Sử dụng logic từ module phân tích KEW để trích xuất các giá trị cực trị (recorded_min/max)
    và giá trị trung bình (avg) cho từng kênh đo.

    Args:
        inps_path: Đường dẫn đến file .KEW.

    Returns:
        dict[str, dict]: Dictionary mapping tên cột với dict {avg, min, max}.
    """
    # Lazy import để gen_word có thể đứng độc lập khi không cần phân tích.
    from modules.kew.analyse_kew import analyse_inps, parse_inps  # type: ignore

    _, df = parse_inps(str(inps_path))
    if df is None:
        return {}
    raw = analyse_inps(df)
    out: dict[str, dict] = {}
    for col, d in raw.items():
        if not isinstance(d, Mapping):
            continue
        avg = d.get("avg")
        rec_min = d.get("recorded_min")
        rec_max = d.get("recorded_max")
        out[col] = {
            "avg": avg,
            "min": rec_min if rec_min is not None else d.get("min"),
            "max": rec_max if rec_max is not None else d.get("max"),
        }
    return out

def _pick(stats: Mapping[str, dict], *keys: str) -> dict:
    """
    Lấy thống kê từ một trong các khóa khớp đầu tiên (không phân biệt hoa thường).
    
    Args:
        stats: Dictionary các thống kê.
        *keys: Các khóa tiềm năng.
        
    Returns:
        dict: Dữ liệu thống kê hoặc dict rỗng.
    """
    upper = {k.upper(): v for k, v in stats.items()}
    for k in keys:
        v = upper.get(k.upper())
        if v:
            return v
    return {}

def _pick_total(stats: Mapping[str, dict], prefix: str, unit_substr: str) -> dict:
    """Tìm kiếm thống kê của cột tổng (không phân biệt pha) dựa trên tiền tố và đơn vị.

    Ví dụ: tìm AVG_PF[_] cho hệ số công suất tổng.

    Args:
        stats: Dictionary chứa toàn bộ thống kê kênh đo.
        prefix: Tiền tố tên cột (VD: 'PF', 'P', 'Q').
        unit_substr: Chuỗi đơn vị nằm trong tên cột (VD: '[W]', '[_]').

    Returns:
        dict: Thống kê tìm được hoặc dict rỗng nếu không khớp.
    """
    for col, d in stats.items():
        if not col.upper().startswith(f"AVG_{prefix.upper()}"):
            continue
        if unit_substr.upper() not in col.upper():
            continue
        if any(ch in col for ch in ("1", "2", "3")):
            continue
        return d
    return {}

def _scale(d: Mapping[str, float] | None, k: float) -> dict:
    """Nhân các giá trị thống kê (min, max, avg) với một hệ số tỉ lệ.

    Args:
        d: Mapping chứa các giá trị số.
        k: Hệ số nhân (VD: 1e-3 để đổi từ W sang kW).

    Returns:
        dict: Dictionary chứa các giá trị đã được nhân tỉ lệ. Trả về dict rỗng nếu d là None.
    """
    if not d:
        return {}
    return {
        "avg": d.get("avg") * k if d.get("avg") is not None else None,
        "min": d.get("min") * k if d.get("min") is not None else None,
        "max": d.get("max") * k if d.get("max") is not None else None,
    }

    """Đánh giá chất lượng điện áp dựa trên độ lệch so với điện áp danh định (±5%).

    Args:
        u_max: Điện áp cực đại đo được.
        u_min: Điện áp cực tiểu đo được.
        u_avg: Điện áp trung bình đo được.
        vref: Điện áp danh định để so sánh (V).

    Returns:
        tuple: (Kết quả "Đạt"/"Chưa đạt"/"Không đạt", độ lệch max, độ lệch min, độ lệch avg).
    """
    if u_max is None or u_min is None or vref <= 0:
        return "—", None, None, None
    dmax = (u_max - vref) / vref * 100
    dmin = (u_min - vref) / vref * 100
    
    dabs_max = max(abs(dmax), abs(dmin))
    dabs_min = min(abs(dmax), abs(dmin))
    
    if u_avg is not None:
        davg = (u_avg - vref) / vref * 100
        avg_ok = abs(davg) <= _V_DEV_LIMIT_PCT
    else:
        avg_ok = True
        
    max_ok = abs(dmax) <= _V_DEV_LIMIT_PCT
    min_ok = abs(dmin) <= _V_DEV_LIMIT_PCT
    
    if not avg_ok:
        res = "Không đạt"
    elif not max_ok or not min_ok:
        res = "Chưa đạt"
    else:
        res = "Đạt"
        
    return res, dabs_max, dabs_min, (dabs_max + dabs_min) / 2

def _eval_pf(pf_max, pf_min, pf_avg) -> str:
    """Đánh giá hệ số công suất (PF) so với ngưỡng 0.9."""
    if pf_avg is None and pf_max is None and pf_min is None:
        return "—"
    
    avg_ok = True if pf_avg is None else abs(pf_avg) >= _PF_LIMIT
    max_ok = True if pf_max is None else abs(pf_max) >= _PF_LIMIT
    min_ok = True if pf_min is None else abs(pf_min) >= _PF_LIMIT
    
    if not avg_ok:
        return "Không đạt"
    elif not max_ok or not min_ok:
        return "Chưa đạt"
    else:
        return "Đạt"

def _eval_thd(values_max: Iterable[float | None], values_avg: Iterable[float | None], limit: float) -> str:
    """Đánh giá tổng biến dạng sóng hài (THD/TDD) so với giới hạn cho phép."""
    max_vals = [v for v in values_max if v is not None]
    avg_vals = [v for v in values_avg if v is not None]
    
    if not max_vals and not avg_vals:
        return "—"
        
    avg_exceeds = any(v >= limit for v in avg_vals)
    max_exceeds = any(v >= limit for v in max_vals)
    
    if avg_exceeds:
        return "Không đạt"
    elif max_exceeds:
        return "Chưa đạt"
    else:
        return "Đạt"

def _eval_unbalance(unb_max: float | None, unb_avg: float | None, limit: float) -> str:
    """Đánh giá độ lệch pha (mất cân bằng) điện áp hoặc dòng điện.

    Args:
        unb_max: Giá trị mất cân bằng cực đại (%).
        unb_avg: Giá trị mất cân bằng trung bình (%).
        limit: Ngưỡng giới hạn cho phép (%).

    Returns:
        str: Kết quả đánh giá ("Đạt", "Chưa đạt", "Không đạt" hoặc "—").
    """
    if unb_avg is not None and unb_max is not None:
        if unb_avg >= limit:
            return "Không đạt"
        elif unb_max >= limit:
            return "Chưa đạt"
        return "Đạt"
    elif unb_max is not None:
        return "Đạt" if unb_max < limit else "Không đạt"
    return "—"

def _fmt_remark_pct(v: float | None, decimals: int = 2) -> str:
    if v is None or v != v:  # NaN
        return "—"
    return f"{v:.{decimals}f}".replace(".", ",")

def _fmt_remark_voltage(v: float | None, decimals: int = 1) -> str:
    if v is None or v != v:
        return "—"
    return f"{v:.{decimals}f}".replace(".", ",")


def _merge_auto_and_excel_notes(auto: str, excel_bits: str) -> str:
    """Ghép đoạn nhận xét sinh tự động với ghi chú bổ sung từ Excel hiện trường.

    Args:
        auto: Nội dung nhận xét tự động.
        excel_bits: Ghi chú thủ công trích xuất từ cột Excel.

    Returns:
        str: Chuỗi văn bản đã ghép, cách nhau bởi 2 dấu xuống dòng.
    """
    bits = (excel_bits or "").strip()
    if not bits:
        return auto
    return f"{auto}\n\n{bits}"

def _device_tdd_limit_from_name(name: str) -> float:
    """Xác định ngưỡng giới hạn TDD dòng điện dựa trên tên/loại thiết bị.

    Args:
        name: Tên thiết bị để nhận diện loại (máy nén, máy ép...).

    Returns:
        float: Ngưỡng giới hạn (%) theo tiêu chuẩn nội bộ (thường là 12.0 hoặc 20.0).
    """
    n = _norm(name)
    if any(k in n for k in ("nén", "nen", "nghiền", "nghien", "máy ép", "may ep", "băng tải", "bang tai")):
        return 12.0
    return 20.0

# ──────────────────────────────────────────────────────────────────────────────
# Sinh nhận xét từ các trường Excel hiện trường
# ──────────────────────────────────────────────────────────────────────────────

_CURRENT_CHAR_MAP: dict[str, str] = {
    "on dinh": "ổn định",
    "on định": "ổn định",
    "ổn định": "ổn định",
    "dao dong nhe": "có sự dao động nhẹ",
    "dao động nhẹ": "có sự dao động nhẹ",
    "bien doi lien tuc": "biến đổi liên tục với biên độ nhỏ",
    "biến đổi liên tục": "biến đổi liên tục với biên độ nhỏ",
    "chu ky load-unload": "biến đổi theo chu kỳ Load/Unload",
    "chu kỳ load-unload": "biến đổi theo chu kỳ Load/Unload",
    "load-unload": "biến đổi theo chu kỳ Load/Unload",
    "load/unload": "biến đổi theo chu kỳ Load/Unload",
}

def _parse_float_field(v: object) -> float | None:
    """Trích xuất giá trị float từ một đối tượng dữ liệu (ô Excel, chuỗi...).

    Hỗ trợ xử lý định dạng dấu phẩy (kiểu Việt Nam) và kiểm tra NaN từ Pandas.

    Args:
        v: Giá trị đầu vào cần chuyển đổi.

    Returns:
        float: Giá trị số thực hoặc None nếu không thể chuyển đổi.
    """
    if v is None:
        return None
    try:
        import pandas as pd
        if isinstance(v, float) and pd.isna(v):
            return None
    except Exception:
        pass
    try:
        s = str(v).strip().replace(",", ".")
        if not s:
            return None
        return float(s)
    except (ValueError, TypeError):
        return None

def _parse_pf_field(v: object) -> float | None:
    """Đọc hệ số cos phi, hỗ trợ cả 0.987 và 987 (chia 1000)."""
    val = _parse_float_field(v)
    if val is None:
        return None
    # Nếu giá trị > 1.0 (VD 987, 850), chia cho 1000 để đưa về dải 0-1.
    if abs(val) > 1.0:
        return val / 1000.0
    return val

def _wave_phrase_from_char(current_char: str | None) -> str:
    """Chuyển giá trị cột current_char → cụm từ tiêu chuẩn."""
    if not current_char:
        return "biến đổi liên tục với biên độ nhỏ"
    key = unicodedata.normalize("NFKC", str(current_char).strip()).lower()
    for k, v in _CURRENT_CHAR_MAP.items():
        if k in key:
            return v
    return str(current_char).strip()

def _pf_phrase(cos_phi: float | None) -> str:
    """Trả về cụm từ mô tả hệ số công suất dựa trên giá trị cosφ."""
    if cos_phi is None or cos_phi != cos_phi:
        return "chưa xác định"
    p = abs(cos_phi)
    if p >= 0.995:
        return "rất cao (cosφ ≈ 1, có thể đã lắp đặt tụ bù)"
    if p >= 0.8:
        return "cao (trên 0,8)"
    if p >= 0.5:
        return "trung bình (dưới 0,8)"
    return "thấp (dưới 0,8)"

def _compose_remarks_from_excel_fields(
    *,
    name: str,
    kind: SectionKind,
    current_char: str | None,
    u_min: float | None,
    u_max: float | None,
    i_max: float | None = None,
    delta_u: float | None,
    delta_i: float | None,
    p_kw: float | None = None,
    cos_phi: float | None,
    thd_max: float | None,
    tdd_max: float | None,
    pdm_kva: float | None = None,
    nominal_voltage: float | None,
) -> str:
    """Sinh nội dung nhận xét tự động dựa trên các thông số kỹ thuật đo được.

    Thuật toán đánh giá (Loi_Dem):
    - Độ lệch điện áp vi phạm: +2 điểm lỗi (lỗi nghiêm trọng).
    - Các vi phạm khác (Cos phi, THD, TDD, lệch dòng): +1 điểm lỗi.
    - Tổng điểm lỗi xác định mức độ đánh giá: tốt, tương đối tốt, hoặc chưa tốt.

    Args:
        name: Tên thiết bị/máy biến áp.
        kind: Loại thiết bị (mba/device/device4).
        current_char: Đặc tính tải (biến đổi/ổn định).
        u_min, u_max, delta_u: Thông số điện áp và độ lệch.
        delta_i: Độ lệch dòng điện.
        p_kw, cos_phi: Công suất và hệ số công suất.
        thd_max, tdd_max: Các chỉ số sóng hài.
        pdm_kva: Công suất định mức (chỉ dùng cho MBA).
        nominal_voltage: Điện áp danh định (V).

    Returns:
        str: Đoạn văn bản nhận xét hoàn chỉnh đã được format.
    """
    vref = float(nominal_voltage) if nominal_voltage and nominal_voltage > 0 else _MBA_NOMINAL_VOLTAGE_V
    wave = _wave_phrase_from_char(current_char)
    name_mid = _format_name_mid(name)

    # ── Định dạng helper ─────────────────────────────────────────────────
    def _pct(v: float | None, d: int = 2) -> str:
        return "—" if v is None else f"{v:.{d}f}".replace(".", ",")

    def _volt(v: float | None, d: int = 1) -> str:
        return "—" if v is None else f"{v:.{d}f}".replace(".", ",")

    # ── δU lệch so với điện áp danh định (từ u_min / u_max) ─────────────
    du_lo: float | None = None
    du_hi: float | None = None
    if u_min is not None and u_max is not None and vref > 0:
        du_lo = (u_min - vref) / vref * 100.0
        du_hi = (u_max - vref) / vref * 100.0

    # ── Thuật toán "Cờ báo lỗi" (Loi_Dem) ───────────────────────────────
    tdd_lim = _TDD_LIMIT_PCT if kind == "mba" else _device_tdd_limit_from_name(name)
    loi_dem = 0

    if cos_phi is not None and abs(cos_phi) < _PF_LIMIT:
        loi_dem += 1
    if du_lo is not None and du_hi is not None:
        if du_lo < -_V_DEV_LIMIT_PCT or du_hi > _V_DEV_LIMIT_PCT:
            loi_dem += 1
    if delta_i is not None and delta_i > 10.0:
        loi_dem += 1
    if thd_max is not None and thd_max > _THDV_LIMIT_PCT:
        loi_dem += 1
    if tdd_max is not None and tdd_max > tdd_lim:
        loi_dem += 1
    if delta_u is not None and delta_u > _V_DEV_LIMIT_PCT:
        loi_dem += 2  # Mất cân bằng điện áp — lỗi nghiêm trọng, cộng 2

    # ── Đánh giá chất lượng tổng quan ────────────────────────────────────
    if kind == "mba":
        quality = "tốt" if loi_dem == 0 else ("tương đối tốt" if loi_dem == 1 else "chưa thực sự tốt")
    else:
        quality = "tốt" if loi_dem == 0 else ("tương đối tốt" if loi_dem == 1 else "chưa tốt")

    # ── Câu Hệ số công suất ───────────────────────────────────────────────
    pf_txt = _pf_phrase(cos_phi)

    # ── Câu Độ lệch pha ΔU / ΔI (Mẫu 1) ─────────────────────────────────
    du_num = float(delta_u) if delta_u is not None else None
    di_num = float(delta_i) if delta_i is not None else None
    du_pass = du_num is not None and du_num <= _V_DEV_LIMIT_PCT
    di_pass = di_num is not None and di_num <= 10.0
    du_s, di_s = _pct(du_num), _pct(di_num)

    if du_num is None and di_num is None:
        unbalance_sent = ""
    elif du_num is None:
        unbalance_sent = (
            f"Độ lệch pha dòng điện ở mức {'thấp' if di_pass else 'cao'} "
            f"(ΔI = {di_s}% {_LT if di_pass else _GT} 10,0%)."
        )
    elif di_num is None:
        unbalance_sent = (
            f"Độ lệch pha điện áp ở mức {'thấp' if du_pass else 'cao'} "
            f"(ΔU = {du_s}% {_LT if du_pass else _GT} 5,0%)."
        )
    elif du_pass and di_pass:
        unbalance_sent = (
            f"Độ lệch pha điện áp và dòng điện đều ở mức thấp "
            f"(ΔU = {du_s}% {_LT} 5,0% {_AMP} ΔI = {di_s}% {_LT} 10,0%)."
        )
    elif du_pass and not di_pass:
        unbalance_sent = (
            f"Độ lệch pha điện áp ở mức thấp (ΔU = {du_s}% {_LT} 5,0%); "
            f"tuy nhiên, độ lệch pha dòng điện vượt mức cho phép (ΔI = {di_s}% {_GT} 10,0%)."
        )
    elif not du_pass and di_pass:
        unbalance_sent = (
            f"Độ lệch dòng điện ở mức thấp (ΔI = {di_s}% {_LT} 10,0%); "
            f"tuy nhiên, độ lệch pha điện áp vượt mức cho phép (ΔU = {du_s}% {_GT} 5,0%)."
        )
    else:
        unbalance_sent = (
            f"Độ lệch pha điện áp và dòng điện đều vượt mức cho phép "
            f"(ΔU = {du_s}% {_GT} 5,0% {_AMP} ΔI = {di_s}% {_GT} 10,0%)."
        )

    # ── Câu Sóng hài THD / TDD (Mẫu 2) ──────────────────────────────────
    lim_s = _pct(tdd_lim, 1)
    th_s, td_s = _pct(thd_max), _pct(tdd_max)
    thd_ok = thd_max is not None and thd_max <= _THDV_LIMIT_PCT
    tdd_ok = tdd_max is not None and tdd_max <= tdd_lim

    if thd_max is None and tdd_max is None:
        harm_sent = ""
    elif thd_ok and tdd_ok:
        harm_sent = (
            f"Tổng biến dạng sóng hài điện áp và dòng điện đều ở mức cho phép "
            f"(THDmax = {th_s}% {_LT} 8,0% {_AMP} TDDmax = {td_s}% {_LT} {lim_s}%)."
        )
    elif thd_ok and not tdd_ok:
        harm_sent = (
            f"Tổng biến dạng sóng hài điện áp ở mức cho phép (THDmax = {th_s}% {_LT} 8,0%); "
            f"tuy nhiên, tổng biến dạng sóng hài dòng điện vượt mức cho phép "
            f"(TDDmax = {td_s}% {_GT} {lim_s}%)."
        )
    elif not thd_ok and tdd_ok:
        harm_sent = (
            f"Tổng biến dạng sóng hài dòng điện ở mức cho phép (TDDmax = {td_s}% {_LT} {lim_s}%); "
            f"tuy nhiên, tổng biến dạng sóng hài điện áp vượt mức cho phép "
            f"(THDmax = {th_s}% {_GT} 8,0%)."
        )
    else:
        harm_sent = (
            f"Tổng biến dạng sóng hài điện áp và tổng biến dạng sóng hài dòng điện "
            f"đều vượt mức cho phép (THDmax = {th_s}% {_GT} 8,0% {_AMP} TDDmax = {td_s}% {_GT} {lim_s}%)."
        )

    # ── MBA: format mới ──────────────────────────────────────────────────────
    # Câu 1: % tải (nếu có đủ dữ liệu)
    # Câu 2: Biểu đồ dòng điện
    # Câu 3: Chất lượng điện + ΔU [+ ΔI nếu cùng mức] + cosφ (một câu)
    # Câu 4: ΔI riêng (chỉ khi ΔI khác mức với ΔU)
    # Câu 5: TDD trong mức (bỏ qua nếu vượt mức)
    # Câu cuối: Dẫn bảng tổng hợp
    if kind == "mba":
        mba_parts: list[str] = []

        # Câu 1 — % công suất tiêu thụ
        if p_kw is not None and cos_phi is not None and abs(cos_phi) > 0.01 and pdm_kva is not None and pdm_kva > 0:
            s_kva = p_kw / abs(cos_phi)
            load_pct = s_kva / pdm_kva * 100.0
            mba_parts.append(
                f"Công suất tiêu thụ của {name_mid} đạt {_pct(load_pct, 2)}% công suất thiết kế."
            )

        # Câu 2 — Biểu đồ dòng điện
        mba_parts.append(
            f"Biểu đồ dòng điện tiêu thụ tại thời điểm đo kiểm {wave}."
        )

        # Câu 3 — Chất lượng điện + ΔU/ΔI + cosφ
        quality_cap = quality[0].upper() + quality[1:]  # Viết hoa chữ đầu
        du_level = "thấp" if (du_num is not None and du_num <= _V_DEV_LIMIT_PCT) else "cao"
        di_level = ("thấp" if di_pass else "cao") if di_num is not None else None

        # Gộp ΔU + ΔI nếu cùng mức, tách câu 4 nếu khác mức
        if di_level is not None and di_level == du_level:
            # Cùng mức → gộp chung vào câu 3
            unbalance_part = f"độ lệch pha điện áp và dòng điện đều ở mức {du_level}"
            di_separate = False
        else:
            # Khác mức hoặc không có ΔI → chỉ nêu ΔU ở câu 3
            unbalance_part = f"độ lệch pha điện áp ở mức {du_level}"
            di_separate = di_level is not None  # Câu 4 riêng nếu có ΔI

        quality_sent = (
            f"Chất lượng điện đo tại {name_mid} ở mức {quality_cap}, "
            f"{unbalance_part}, "
            f"hệ số công suất cosφ ở mức {pf_txt}."
        )
        mba_parts.append(quality_sent)

        # Câu 4 — ΔI riêng (chỉ khi khác mức ΔU)
        if di_separate:
            mba_parts.append(
                f"Độ lệch pha dòng điện ở mức {di_level} (ΔI = {di_s}%)."
            )

        # Câu 5 — TDD dòng điện trong mức (bỏ qua nếu TDD vượt mức)
        if tdd_ok:
            mba_parts.append(
                f"Tổng biến dạng sóng hài dòng điện ở mức cho phép "
                f"(TDDmax = {td_s}% {_LT} {lim_s}%)."
            )

        # Câu cuối — Dẫn bảng
        mba_parts.append(
            f"Dưới đây là bảng tổng hợp thông số hoạt động của {name_mid}:"
        )

        return " ".join(mba_parts)

    # ── Device: 6 câu theo thứ tự spec ───────────────────────────────────
    # Câu 4: Điện áp dao động + δU so danh định
    umin_s, umax_s = _volt(u_min), _volt(u_max)
    dlo_s, dhi_s = _pct(du_lo), _pct(du_hi)
    if du_lo is not None and du_hi is not None:
        both_in = abs(du_lo) <= _V_DEV_LIMIT_PCT and abs(du_hi) <= _V_DEV_LIMIT_PCT
        verdict = "đạt tiêu chuẩn (-5,0% ≤ δ ≤ 5,0%)" if both_in else "vượt giới hạn cho phép (-5,0% ≤ δ ≤ 5,0%)"
        volt_sent = (
            f"Điện áp dao động từ {umin_s} ÷ {umax_s} V, "
            f"độ lệch chuẩn của điện áp δU (= {dlo_s}% ÷ {dhi_s}%), {verdict}."
        )
    elif u_min is not None and u_max is not None:
        volt_sent = f"Điện áp dao động từ {umin_s} ÷ {umax_s} V."
    else:
        volt_sent = ""

    parts: list[str] = [
        f"Chất lượng điện cấp cho {name_mid} ở mức {quality}.",
        f"Biểu đồ dòng điện tiêu thụ {wave} trong thời gian đo kiểm.",
        f"Hệ số công suất cosφ ở mức {pf_txt}.",
    ]
    if volt_sent:
        parts.append(volt_sent)
    if unbalance_sent:
        parts.append(unbalance_sent)
    if harm_sent:
        parts.append(harm_sent)
    return " ".join(parts)

def _resolve_remarks_field(
    *,
    kind: SectionKind,
    folder: Path,
    name: str,
    user_remarks: str,
    nominal_voltage: float | None,
    excel_params: dict | None = None,
) -> str:
    """Quyết định nội dung nhận xét dựa trên dữ liệu Excel hiện trường hoặc tự động.

    Args:
        kind: Loại thiết bị (mba/device/device4).
        folder: Thư mục chứa dữ liệu thiết bị.
        name: Tên thiết bị.
        user_remarks: Ghi chú thủ công từ người dùng (nếu có).
        nominal_voltage: Điện áp danh định (V).
        excel_params: Các thông số đo kiểm từ Excel (P, Q, S, Cosphi, Harmonic...).

    Returns:
        str: Chuỗi văn bản nhận xét hoàn chỉnh đã được format.
    """
    raw = (user_remarks or "").strip()

    params = excel_params or {}
    current_char = params.get("current_char")
    u_min = _parse_float_field(params.get("u_min"))
    u_max = _parse_float_field(params.get("u_max"))
    i_max = _parse_float_field(params.get("i_max"))
    delta_u = _parse_float_field(params.get("delta_u"))
    delta_i = _parse_float_field(params.get("delta_i"))
    p_kw = _parse_float_field(params.get("p"))
    cos_phi = _parse_pf_field(params.get("cos_phi"))
    thd_max = _parse_float_field(params.get("thd"))
    tdd_max_v = _parse_float_field(params.get("tdd"))
    pdm_kva = _parse_float_field(params.get("pdm"))

    has_data = any(v is not None for v in [
        current_char, u_min, u_max, delta_u, delta_i, cos_phi, thd_max, tdd_max_v
    ])
    if not has_data:
        return raw

    auto = _compose_remarks_from_excel_fields(
        name=name,
        kind=kind,
        current_char=current_char,
        u_min=u_min,
        u_max=u_max,
        i_max=i_max,
        delta_u=delta_u,
        delta_i=delta_i,
        p_kw=p_kw,
        cos_phi=cos_phi,
        thd_max=thd_max,
        tdd_max=tdd_max_v,
        pdm_kva=pdm_kva,
        nominal_voltage=nominal_voltage,
    )
    return _merge_auto_and_excel_notes(auto, raw)

# ════════════════════════════════════════════════════════════════════
#                       Image discovery helpers
# ════════════════════════════════════════════════════════════════════

def list_bmp_in_folder(folder: str | Path) -> list[Path]:
    """Tìm kiếm và sắp xếp các file ảnh PS-SDxxxx.BMP trong thư mục.

    Args:
        folder: Thư mục cần quét ảnh.

    Returns:
        list[Path]: Danh sách các đường dẫn file BMP, sắp xếp theo số hiệu ảnh.
    """
    p = Path(folder)
    if not p.is_dir():
        return []
    bmps: list[tuple[int, Path]] = []
    for f in p.iterdir():
        if not f.is_file():
            continue
        m = _BMP_RE.match(f.name)
        if m:
            bmps.append((int(m.group(1)), f))
    bmps.sort(key=lambda x: x[0])
    # Fallback: mọi file .bmp nếu không theo định dạng đặt tên chuẩn
    if not bmps:
        extra: list[tuple[int, Path]] = []
        for i, f in enumerate(sorted(p.glob("*.BMP")) + sorted(p.glob("*.bmp"))):
            if not f.is_file():
                continue
            extra.append((i, f))
        bmps = extra
    return [pp for _, pp in bmps]

def find_overview_png(folder: str | Path) -> Path | None:
    """Tìm file ảnh tổng quan (mặc định là a.png) trong thư mục thiết bị.

    Args:
        folder: Thư mục cần tìm kiếm.

    Returns:
        Path | None: Đường dẫn file ảnh nếu tìm thấy, ngược lại là None.
    """
    p = Path(folder)
    if not p.is_dir():
        return None
    for f in p.iterdir():
        if not f.is_file():
            continue
        if f.suffix.lower() == ".png" and f.stem.lower() == "a":
            return f
    return None

def _require_overview_png_and_bmps(folder: str | Path) -> tuple[Path, list[Path]]:
    """Trả về ``(a.png, danh_sách_PS-SD_BMP)`` hoặc ném lỗi nếu thiếu."""
    p = Path(folder)
    overview = find_overview_png(p)
    if overview is None:
        raise FileNotFoundError(
            f"Thiếu file a.png (ảnh tổng quan) trong {p!s}. "
            "Vui lòng đặt a.png trong thư mục thiết bị."
        )
    bmps = list_bmp_in_folder(p)
    if not bmps:
        raise FileNotFoundError(f"Không tìm thấy PS-SDxxxx.BMP trong {p!s}.")
    return overview, bmps

def _take(lst: list[Path], idx: int, fallback: Path | None) -> Path | None:
    return lst[idx] if 0 <= idx < len(lst) else fallback

def auto_pick_mba_images(folder: str | Path) -> dict[str, str]:
    """Tự động chọn bộ ảnh phù hợp cho template báo cáo MBA.

    Yêu cầu:
    - 1 file 'a.png' cho ảnh tổng quan (imga).
    - Các file 'PS-SDxxxx.BMP' cho ảnh thông số (img1, img2, img4, img6).

    Args:
        folder: Thư mục chứa ảnh của máy biến áp.

    Returns:
        dict: Mapping các khóa template ảnh với đường dẫn file thực tế.
    """
    overview, bmps = _require_overview_png_and_bmps(folder)
    fb = bmps[0]
    imga = str(overview)
    return {
        "imga": imga,
        "img1": str(_take(bmps, 0, fb)),
        "img2": str(_take(bmps, 1, fb)),
        "img4": str(_take(bmps, 2, fb)),
        "img6": str(_take(bmps, 3, fb)),
    }

def auto_pick_device_images(folder: str | Path) -> dict[str, str]:
    """Chọn ảnh cho template device (``imga`` + ``img1``…``img6``).

    * ``imga``: **bắt buộc** file ``a.png`` trong thư mục.
    * ``img1``…``img6``: ``PS-SDxxxx.BMP`` theo thứ tự số (vd 641–646 → ``img1``–``img6``).
    """
    overview, bmps = _require_overview_png_and_bmps(folder)
    fb = bmps[0]
    imga = str(overview)
    return {
        "imga": imga,
        **{f"img{i}": str(_take(bmps, i - 1, fb)) for i in range(1, 7)},
    }

# ════════════════════════════════════════════════════════════════════
#                  Kwargs builders từ thư mục thiết bị
# ════════════════════════════════════════════════════════════════════

def mba_kwargs_from_inps(
    inps_path: str | Path | None,
    *,
    name: str,
    imga: str,
    img1: str,
    img2: str,
    img4: str,
    img6: str,
    cap_fig_mba: str | None = None,
    remarks_mba: str = "",
    cap_tab_mba: str | None = None,
) -> dict:
    """Đọc dữ liệu từ file INPSxxxx.KEW và chuẩn bị bộ tham số (kwargs) để render template MBA.

    Args:
        inps_path: Đường dẫn tới file INPSxxxx.KEW. Nếu None, các giá trị sẽ là "—".
        name: Tên máy biến áp.
        imga, img1, img2, img4, img6: Đường dẫn tới các file ảnh đo kiểm.
        cap_fig_mba: Chú thích hình (tuỳ chọn).
        remarks_mba: Nhận xét văn bản bổ sung (tuỳ chọn).
        cap_tab_mba: Chú thích bảng thông số (tuỳ chọn).

    Returns:
        dict: Context dữ liệu đầy đủ cho template MBA.
    """
    if inps_path is None:
        stats: dict[str, dict] = {}
    else:
        stats = _parse_inps(inps_path)

    # ─── Điện áp dây U12, U23, U31 (ưu tiên AVG_VLi[V]) ────────────
    u12 = _pick(stats, "AVG_VL1[V]", "AVG_V12[V]")
    u23 = _pick(stats, "AVG_VL2[V]", "AVG_V23[V]")
    u31 = _pick(stats, "AVG_VL3[V]", "AVG_V31[V]")

    # Nếu thiết bị chỉ ghi điện áp pha → quy đổi sang dây bằng √3.
    if not u12:
        v1 = _pick(stats, "AVG_V1[V]")
        v2 = _pick(stats, "AVG_V2[V]")
        v3 = _pick(stats, "AVG_V3[V]")
        k = 3 ** 0.5
        u12, u23, u31 = _scale(v1, k), _scale(v2, k), _scale(v3, k)

    vref = _MBA_NOMINAL_VOLTAGE_V
    u12eval, _, _, _ = _eval_voltage(u12.get("max"), u12.get("min"), u12.get("avg"), vref)

    # ─── Dòng điện I1, I2, I3 ─────────────────────────────────────
    i1 = _pick(stats, "AVG_A1[A]")
    i2 = _pick(stats, "AVG_A2[A]")
    i3 = _pick(stats, "AVG_A3[A]")

    # ─── Độ lệch pha điện áp / dòng (%) ───────────────────────────
    uv_unb = _pick(stats, "AVG_UV[%]", "AVG_VUNB[%]")
    ua_unb = _pick(stats, "AVG_UA[%]", "AVG_AUNB[%]")
    dueval = _eval_unbalance(uv_unb.get("max"), uv_unb.get("avg"), _V_DEV_LIMIT_PCT)

    # ─── Hệ số công suất ─────────────────────────────────────────
    pf = _pick_total(stats, "PF", "[_]") or _pick(stats, "AVG_PF[_]")
    pfeval = _eval_pf(pf.get("max"), pf.get("min"), pf.get("avg"))

    # ─── P, Q, S (đổi sang kW/kvar/kVA) ──────────────────────────
    p_total = _pick_total(stats, "P", "[W]") or _pick(stats, "AVG_P[W]")
    q_total = _pick_total(stats, "Q", "[var]") or _pick(stats, "AVG_Q[var]")
    s_total = _pick_total(stats, "S", "[VA]") or _pick(stats, "AVG_S[VA]")
    p_k = _scale(p_total, 1e-3)
    q_k = _scale(q_total, 1e-3)
    s_k = _scale(s_total, 1e-3)

    # ─── THD V & TDD I theo pha ─────────────────────────────────
    thd1 = _pick(stats, "AVG_THDVR1[%]", "AVG_VTHD1[%]")
    thd2 = _pick(stats, "AVG_THDVR2[%]", "AVG_VTHD2[%]")
    thd3 = _pick(stats, "AVG_THDVR3[%]", "AVG_VTHD3[%]")
    tdd1 = _pick(stats, "AVG_THDAR1[%]", "AVG_ATHD1[%]")
    tdd2 = _pick(stats, "AVG_THDAR2[%]", "AVG_ATHD2[%]")
    tdd3 = _pick(stats, "AVG_THDAR3[%]", "AVG_ATHD3[%]")

    thdeval = _eval_thd(
        [thd1.get("max"), thd2.get("max"), thd3.get("max")],
        [thd1.get("avg"), thd2.get("avg"), thd3.get("avg")],
        _THDV_LIMIT_PCT
    )
    tddeval = _eval_thd(
        [tdd1.get("max"), tdd2.get("max"), tdd3.get("max")],
        [tdd1.get("avg"), tdd2.get("avg"), tdd3.get("avg")],
        _TDD_LIMIT_PCT
    )

    u12max, u12min, u12avg = _tri(u12, 1)
    u23max, u23min, u23avg = _tri(u23, 1)
    u31max, u31min, u31avg = _tri(u31, 1)
    i1max, i1min, i1avg = _tri(i1, 0)
    i2max, i2min, i2avg = _tri(i2, 0)
    i3max, i3min, i3avg = _tri(i3, 0)
    dumax, dumin, duavg = _tri(uv_unb, 3)
    dimax, dimin, diavg = _tri(ua_unb, 3)
    pfmax, pfmin, pfavg = _tri(pf, 3)
    pmax, pmin, pavg = _tri(p_k, 1)
    qmax, qmin, qavg = _tri(q_k, 1)
    smax, smin, savg = _tri(s_k, 1)
    thd1max, thd1min, thd1avg = _tri(thd1, 2)
    thd2max, thd2min, thd2avg = _tri(thd2, 2)
    thd3max, thd3min, thd3avg = _tri(thd3, 2)
    tdd1max, tdd1min, tdd1avg = _tri(tdd1, 2)
    tdd2max, tdd2min, tdd2avg = _tri(tdd2, 2)
    tdd3max, tdd3min, tdd3avg = _tri(tdd3, 2)

    nm = _format_name_mid(name)
    return {
        "name": name,
        "imga": imga, "img1": img1, "img2": img2, "img4": img4, "img6": img6,
        "cap_fig_mba": cap_fig_mba if cap_fig_mba is not None else f"Kết quả đo chất lượng điện {nm}",
        "remarks_mba": remarks_mba,
        "cap_tab_mba": cap_tab_mba if cap_tab_mba is not None else f"Thông số hoạt động của {nm}",
        "u12max": u12max, "u12min": u12min, "u12avg": u12avg, "u12eval": u12eval,
        "u23max": u23max, "u23min": u23min, "u23avg": u23avg,
        "u31max": u31max, "u31min": u31min, "u31avg": u31avg,
        "i1max": i1max, "i1min": i1min, "i1avg": i1avg,
        "i2max": i2max, "i2min": i2min, "i2avg": i2avg,
        "i3max": i3max, "i3min": i3min, "i3avg": i3avg,
        "dumax": dumax, "dumin": dumin, "duavg": duavg, "dueval": dueval,
        "dimax": dimax, "dimin": dimin, "diavg": diavg,
        "pfmax": pfmax, "pfmin": pfmin, "pfavg": pfavg, "pfeval": pfeval,
        "pmax": pmax, "pmin": pmin, "pavg": pavg,
        "qmax": qmax, "qmin": qmin, "qavg": qavg,
        "smax": smax, "smin": smin, "savg": savg,
        "thd1max": thd1max, "thd1min": thd1min, "thd1avg": thd1avg, "thdeval": thdeval,
        "thd2max": thd2max, "thd2min": thd2min, "thd2avg": thd2avg,
        "thd3max": thd3max, "thd3min": thd3min, "thd3avg": thd3avg,
        "tdd1max": tdd1max, "tdd1min": tdd1min, "tdd1avg": tdd1avg, "tddeval": tddeval,
        "tdd2max": tdd2max, "tdd2min": tdd2min, "tdd2avg": tdd2avg,
        "tdd3max": tdd3max, "tdd3min": tdd3min, "tdd3avg": tdd3avg,
    }

def mba_kwargs_from_folder(
    folder: str | Path,
    *,
    name: str,
    cap_fig_mba: str | None = None,
    remarks_mba: str = "",
    cap_tab_mba: str | None = None,
    excel_params: dict | None = None,
) -> dict:
    """Tự động chọn ảnh và dữ liệu từ thư mục để dựng tham số cho template MBA.

    Quy trình:
    1. Tìm file dữ liệu gốc (INPSxxxx.KEW) trong thư mục bằng tiền tố "INPS".
    2. Tự động chọn các file ảnh đo kiểm (biểu đồ dòng, sóng hài...).
    3. Sinh nội dung nhận xét dựa trên dữ liệu Excel hiện trường.

    Args:
        folder: Thư mục chứa dữ liệu của máy biến áp.
        name: Tên hiển thị của MBA.
        cap_fig_mba: Chú thích hình ảnh.
        remarks_mba: Nhận xét bổ sung.
        cap_tab_mba: Chú thích bảng.
        excel_params: Các tham số đo lường bổ sung từ Excel (nếu có).

    Returns:
        dict: Context dữ liệu sẵn sàng để truyền vào hàm mba().
    """
    folder = Path(folder)
    from modules.kew.analyse_kew import find_file  # type: ignore

    inps_path = find_file(str(folder), "INPS")
    images = auto_pick_mba_images(folder)
    kw = mba_kwargs_from_inps(
        inps_path,
        name=name,
        cap_fig_mba=cap_fig_mba,
        remarks_mba="",
        cap_tab_mba=cap_tab_mba,
        **images,
    )
    kw["remarks_mba"] = _resolve_remarks_field(
        kind="mba",
        folder=folder,
        name=name,
        user_remarks=remarks_mba,
        nominal_voltage=None,
        excel_params=excel_params,
    )
    return kw

def device_kwargs_from_folder(
    folder: str | Path,
    *,
    name: str,
    cap_device: str | None = None,
    remarks_device: str = "",
    nominal_voltage: float | None = None,
    excel_params: dict | None = None,
) -> dict:
    """Tự động chọn ảnh và dữ liệu từ thư mục để dựng tham số cho template thiết bị đo kiểm.

    Args:
        folder: Đường dẫn thư mục chứa dữ liệu thiết bị.
        name: Tên hiển thị của thiết bị.
        cap_device: Chú thích cho hình ảnh và bảng dữ liệu (tuỳ chọn).
        remarks_device: Ghi chú bổ sung từ người dùng (tuỳ chọn).
        nominal_voltage: Điện áp danh định để tính toán độ lệch (V).
        excel_params: Các tham số đo kiểm trích xuất từ Excel.

    Returns:
        dict: Tập hợp các tham số (kwargs) sẵn sàng để truyền vào hàm :func:`device`.
    """
    images = auto_pick_device_images(folder)
    nm = _format_name_mid(name)
    return {
        "name": name,
        "cap_device": cap_device if cap_device is not None else f"Kết quả đo chất lượng điện {nm}",
        "remarks_device": _resolve_remarks_field(
            kind="device",
            folder=Path(folder),
            name=name,
            user_remarks=remarks_device,
            nominal_voltage=nominal_voltage,
            excel_params=excel_params,
        ),
        **images,
    }

# ════════════════════════════════════════════════════════════════════
#                         Render & merge
# ════════════════════════════════════════════════════════════════════

def merge_rendered_docx(
    docx_paths: Sequence[str | Path],
    output_path: str | Path,
) -> Path:
    """Ghép nối nhiều file Word thành một file duy nhất, bảo toàn định dạng và media.

    Args:
        docx_paths: Danh sách đường dẫn các file .docx cần ghép.
        output_path: Đường dẫn lưu file kết quả.

    Returns:
        Path: Đường dẫn file kết quả sau khi ghép.
    """
    paths = [Path(p) for p in docx_paths]
    if not paths:
        raise ValueError("merge_rendered_docx: cần ít nhất một đường dẫn .docx.")
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    if len(paths) == 1:
        copy2(paths[0], out)
        return out
    composer = Composer(Document(str(paths[0])))
    for p in paths[1:]:
        composer.append(Document(str(p)))
    composer.save(str(out))
    return out

def merge_mba_device_docx(
    output_path: str | Path,
    *,
    mba_template: str | Path,
    device_template: str | Path,
    device4_template: str | Path | None = None,
    totalmba_template: str | Path | None = None,
    mba_sections: list[dict] | None = None,
    device_sections: list[dict] | None = None,
    sections: Sequence[tuple[SectionKind, dict]] | None = None,
) -> Path:
    """Render từng MBA / device rồi ghép thành ``output_path``.

    * ``sections``: danh sách ``("mba", kwargs)`` / ``("device", kwargs)`` theo
      đúng thứ tự cần xuất hiện trong file cuối.
    * Nếu ``sections=None``: lần lượt toàn bộ ``mba_sections`` rồi ``device_sections``.
    """
    mba_sections = mba_sections or []
    device_sections = device_sections or []
    mba_tpl, device_tpl = Path(mba_template), Path(device_template)
    device4_tpl = Path(device4_template) if device4_template else Path(DEFAULT_DEVICE4_TEMPLATE)
    totalmba_tpl = Path(totalmba_template) if totalmba_template else Path(DEFAULT_TOTALMBA_TEMPLATE)

    if sections is not None:
        work: list[tuple[SectionKind, dict]] = list(sections)
    else:
        work = [("mba", s) for s in mba_sections] + [("device", s) for s in device_sections]

    if not work:
        raise ValueError(
            "merge_mba_device_docx: cần sections hoặc mba_sections/device_sections không rỗng."
        )

    kind_render = {
        "mba": (mba_tpl, mba),
        "device": (device_tpl, device),
        "device4": (device4_tpl, device),
        "total_mba": (totalmba_tpl, total_mba),
    }

    with TemporaryDirectory() as td:
        tmp = Path(td)
        rendered: list[Path] = []
        for i, (kind, spec) in enumerate(work):
            try:
                tpl_path, render_fn = kind_render[kind]
            except KeyError as e:
                raise ValueError(
                    f"Loại section không hợp lệ: {kind!r} (chỉ 'mba', 'device', 'device4')."
                ) from e
            tpl = DocxTemplate(str(tpl_path))
            tpl.render(render_fn(tpl, **spec), autoescape=True)
            path = tmp / f"{kind}_{i}.docx"
            tpl.save(str(path))
            if i < len(work) - 1:
                doc_pb = Document(str(path))
                doc_pb.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                doc_pb.save(str(path))
            rendered.append(path)
        return merge_rendered_docx(rendered, output_path)

# ════════════════════════════════════════════════════════════════════
#                Pipeline cho luồng "Xử lý file sơ bộ"
# ════════════════════════════════════════════════════════════════════

_MBA_NAME_RE = re.compile(r"^(MBA|TR|TBA|T\d|MBT)\b|MÁY BIẾN ÁP|BIẾN ÁP", re.IGNORECASE)

def _guess_kind(name: str) -> SectionKind:
    """Đoán loại template từ tên thiết bị (chỉ khi không có cột ``type`` từ Excel)."""
    return "mba" if _MBA_NAME_RE.search(name or "") else "device"

def _resolve_word_section_kind(
    spec: Mapping,
    *,
    name: str,
    default_kind: SectionKind | None,
) -> SectionKind:
    """Chọn ``mba`` / ``device`` / ``device4`` cho báo cáo Word.

    * Nếu mục có khóa ``kind`` (luồng ZIP + Excel): ``mba`` khi cột type nhận
      diện MBA; ``device4`` khi type là ``"4"``; ô trống / không nhận diện /
      ``device`` → ``device`` (không đoán theo tên).
    * Nếu không có khóa ``kind`` (quét thư mục thuần): ``default_kind`` hoặc
      :func:`_guess_kind`.
    """
    if "kind" in spec:
        raw = spec["kind"]
        if raw is not None and not (isinstance(raw, str) and not str(raw).strip()):
            nk = _norm_kind(raw)
            if nk == "mba":
                return "mba"
            if nk == "device4":
                return "device4"
            if nk == "device":
                return "device"
        return "device"
    if default_kind in ("mba", "device", "device4"):
        return default_kind
    return _guess_kind(name)

def build_field_word_report(
    project_root: str | Path,
    output_path: str | Path,
    *,
    mba_template: str | Path,
    device_template: str | Path,
    device4_template: str | Path | None = None,
    totalmba_template: str | Path | None = None,
    devices: Sequence[Mapping] | None = None,
    default_kind: SectionKind | None = None,
    chapter_filter: Literal["all", "chapter4", "chapter5"] = "all",
) -> tuple[Path, list[str]]:
    """Quét thư mục gốc và sinh file báo cáo Word tổng hợp (Chương 4, Chương 5 hoặc cả hai).

    Quy trình xử lý:
    1. Quét toàn bộ các thư mục con trong dự án và trích xuất dữ liệu đo kiểm.
    2. Phân loại thiết bị dựa trên metadata (MBA, Device, Device4).
    3. Lọc danh sách thiết bị dựa trên ``chapter_filter``.
    4. Tổng hợp bảng thông số ``ds_mba`` cho Chương 5.
    5. Render từng section bằng template tương ứng và ghép thành file cuối cùng.

    Args:
        project_root: Thư mục chứa các thư mục con của từng thiết bị (VD: Project_Output).
        output_path: Đường dẫn lưu file .docx kết quả.
        mba_template: Template cho MBA.
        device_template: Template cho thiết bị thường.
        device4_template: Template cho thiết bị loại 4.
        totalmba_template: Template cho bảng danh sách MBA.
        devices: Danh sách metadata (tên, thư mục, tham số Excel) đã chuẩn bị sẵn.
        default_kind: Loại thiết bị mặc định nếu không xác định được.
        chapter_filter: Chế độ lọc chương ("all", "chapter4", "chapter5").

    Returns:
        tuple[Path, list[str]]: (Đường dẫn file Word, danh sách các cảnh báo trong quá trình sinh).
    """
    root = Path(project_root)
    if not root.is_dir():
        raise FileNotFoundError(f"Không tìm thấy thư mục Project_Output: {root}")

    if devices is None:
        devices = [
            {"name": _nfc(d.name), "folder": d}
            for d in sorted(root.iterdir())
            if d.is_dir() and not d.name.startswith(".") and d.name != "__MACOSX"
        ]
    if not devices:
        raise ValueError("Không có thiết bị nào để dựng báo cáo Word.")

    warnings: list[str] = []

    ds_mba: list[dict] = []
    mba_count = 0
    for spec in devices:
        name = _nfc(str(spec.get("name") or "").strip())
        if not name:
            continue
        kind = _resolve_word_section_kind(spec, name=name, default_kind=default_kind)
        # ds_mba chỉ xuất hiện trong chương 5 (chapter5 hoặc all)
        if kind == "mba" and chapter_filter != "chapter4":
            mba_count += 1
            ep = spec.get("excel_params") or {}

            def _fmt(v, dec=2):
                if v is None:
                    return "—"
                try:
                    import pandas as _pd
                    if isinstance(v, float) and _pd.isna(v):
                        return "—"
                except Exception:
                    pass
                try:
                    x = float(str(v).strip().replace(",", "."))
                    if x != x:
                        return "—"
                    return f"{x:.{dec}f}".replace(".", ",")
                except Exception:
                    return str(v).strip() if str(v).strip() else "—"

            ds_mba.append({
                "tt": mba_count,
                "ten": name,
                "pdm": _fmt(ep.get("pdm"), 0),
                "p": _fmt(ep.get("p"), 0),
                "pf": _fmt(_parse_pf_field(ep.get("cos_phi")), 3),
            })

    sections: list[tuple[SectionKind, dict]] = []
    for spec in devices:
        name = _nfc(str(spec.get("name") or "").strip())
        folder_raw = spec.get("folder")
        if not name or not folder_raw:
            warnings.append(f"Bỏ qua mục thiếu name/folder: {spec!r}.")
            continue
        folder = Path(folder_raw)
        if not folder.is_absolute():
            folder = root / folder
        if not folder.is_dir():
            warnings.append(f"«{name}»: không tìm thấy thư mục {folder}.")
            continue

        kind = _resolve_word_section_kind(spec, name=name, default_kind=default_kind)

        # Lọc theo chapter_filter
        if chapter_filter == "chapter4" and kind != "device4":
            continue
        if chapter_filter == "chapter5" and kind == "device4":
            continue

        remarks = str(spec.get("remarks") or "")
        nom_raw = spec.get("nominal_voltage")
        nom_v: float | None = None
        if isinstance(nom_raw, (int, float)) and not isinstance(nom_raw, bool):
            if not (isinstance(nom_raw, float) and nom_raw != nom_raw):
                nom_v = float(nom_raw)
        excel_params = spec.get("excel_params") or {}

        try:
            if kind == "mba":
                kwargs = mba_kwargs_from_folder(
                    folder, name=name, remarks_mba=remarks,
                    excel_params=excel_params,
                )
                kwargs["ds_mba"] = ds_mba
            else:  # "device" or "device4"
                kwargs = device_kwargs_from_folder(
                    folder,
                    name=name,
                    remarks_device=remarks,
                    nominal_voltage=nom_v,
                    excel_params=excel_params,
                )
            sections.append((kind, kwargs))
        except FileNotFoundError as e:
            warnings.append(f"«{name}»: {e}.")
        except Exception as e:
            warnings.append(f"«{name}»: lỗi dựng section ({e}).")

    # Chỉ chèn bảng tổng hợp MBA trong chương 5 (hoặc all)
    if ds_mba and chapter_filter != "chapter4":
        insert_idx = len(sections)
        for i, (k, _) in enumerate(sections):
            if k in ("device", "device4"):
                insert_idx = i
                break
        sections.insert(insert_idx, ("total_mba", {"ds_mba": ds_mba}))

    if not sections:
        raise RuntimeError("Không dựng được section nào: " + "; ".join(warnings))

    out = Path(output_path)
    merge_mba_device_docx(
        out,
        mba_template=mba_template,
        device_template=device_template,
        device4_template=device4_template,
        totalmba_template=totalmba_template,
        sections=sections,
    )
    return out, warnings

# ════════════════════════════════════════════════════════════════════
#   Đọc Excel metadata (Loại / U định mức / Nhận xét) cho tab Word
# ════════════════════════════════════════════════════════════════════

def _norm(s: object) -> str:
    if s is None:
        return ""
    t = unicodedata.normalize("NFKC", str(s)).strip().lower()
    return re.sub(r"\s+", " ", t)

def _nfc(s: object) -> str:
    """Chuẩn hoá tên hiển thị (tránh ký tự tổ hợp kiểu Mac / Excel lệch dạng)."""
    if s is None:
        return ""
    return unicodedata.normalize("NFC", str(s).strip())

_DEVICE4_KIND_LABELS = frozenset({"4", "device4"})

def _norm_kind(value: object) -> SectionKind | None:
    raw = value
    if raw is None or (isinstance(raw, float) and raw != raw):
        return None
    k = _norm(raw)
    if not k:
        return None
    if k == "mba":
        return "mba"
    if k in _DEVICE4_KIND_LABELS:
        return "device4"
    if k in _DEVICE_KIND_LABELS:
        return "device"
    return None

def _metadata_first_hit(metadata: dict[str, dict], folder_name: str) -> dict | None:
    for v in (
        folder_name,
        unicodedata.normalize("NFC", folder_name),
        unicodedata.normalize("NFD", folder_name),
    ):
        hit = metadata.get(_norm(v))
        if hit:
            return hit
    return None

def _metadata_keys_for_excel_name(name: str) -> list[str]:
    """Các khóa tra cứu metadata cho một dòng Excel (tên gốc + tên thư mục sau sanitize)."""
    keys: list[str] = [_norm(name)]
    try:
        from modules.kew.organize_field_zip import sanitize_device_folder

        san = sanitize_device_folder(name)
    except (ValueError, ImportError):
        san = ""
    if san:
        for v in (san, unicodedata.normalize("NFC", san), unicodedata.normalize("NFD", san)):
            nk = _norm(v)
            if nk not in keys:
                keys.append(nk)
    return keys

def _lookup_device_metadata(metadata: dict[str, dict], folder_name: str) -> dict:
    """Ghép dòng Excel với thư mục ``Project_Output/<tên>/`` (NFC/NFD, hậu tố ``_2``…)."""
    if not metadata:
        return {}
    hit = _metadata_first_hit(metadata, folder_name)
    if hit:
        return hit
    m = re.fullmatch(r"(.+)_([1-9]\d*)$", folder_name)
    if m:
        base = m.group(1)
        hit = _metadata_first_hit(metadata, base)
        if hit:
            return hit
    m2 = re.match(r"^\s*([Ss]\d{1,4})\s*[-–—]?\s*", folder_name)
    if m2:
        code = m2.group(1).upper()
        matches: list[dict] = []
        seen: set[int] = set()
        for ent in metadata.values():
            i = id(ent)
            if i in seen:
                continue
            seen.add(i)
            dn = str(ent.get("name") or "")
            if re.match(rf"^\s*{re.escape(code)}\s*[-–—]?\s*", dn, re.IGNORECASE):
                matches.append(ent)
        if len(matches) == 1:
            return matches[0]
    return {}

def _excel_stt(value: object) -> int | None:
    """Trích xuất và chuẩn hóa số thứ tự (STT) từ ô dữ liệu Excel.

    Args:
        value: Giá trị thô từ ô Excel.

    Returns:
        int | None: Số nguyên đại diện cho STT, hoặc None nếu không hợp lệ.
    """
    if value is None:
        return None
    if isinstance(value, float) and value != value:  # NaN
        return None
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        if isinstance(value, float) and value.is_integer():
            return int(value)
        if isinstance(value, int):
            return int(value)
    s = str(value).strip()
    if not s:
        return None
    digits = re.sub(r"\D", "", s)
    if not digits:
        return None
    return int(digits)

def _cell_text(row: object, col: str | None) -> str:
    """Trích xuất chuỗi văn bản sạch từ một ô dữ liệu trong hàng (Series) của Pandas.

    Args:
        row: Hàng dữ liệu (Series).
        col: Tên cột cần lấy dữ liệu.

    Returns:
        str: Chuỗi văn bản đã loại bỏ khoảng trắng hoặc chuỗi rỗng nếu ô trống/NaN.
    """
    if col is None:
        return ""
    try:
        import pandas as pd

        v = row[col]
        if v is None or (isinstance(v, float) and pd.isna(v)):
            return ""
    except Exception:
        return ""
    return str(v).strip()

def read_device_metadata_from_excel(
    excel_path: str | Path,
) -> dict[str, dict]:
    """Đọc dữ liệu metadata và tham số đo kiểm từ file Excel hiện trường.

    Args:
        excel_path: Đường dẫn đến file Excel (.xlsx).

    Returns:
        dict[str, dict]: Mapping từ tên thiết bị (đã chuẩn hóa) sang thông tin chi tiết
            (stt, loại, điện áp danh định, tham số đo lường...).
    """
    try:
        import pandas as pd

        from modules.kew.organize_field_zip import resolve_field_excel_column_map
    except ImportError:
        return {}
    try:
        df = pd.read_excel(str(excel_path), header=0, engine="openpyxl")
    except Exception:
        return {}
    if df.empty:
        return {}

    try:
        cm = resolve_field_excel_column_map(df)
    except ValueError:
        return {}

    name_col = cm["name"]
    kind_col = cm["type"]

    out: dict[str, dict] = {}
    for _, row in df.iterrows():
        raw_name = row[name_col]
        if raw_name is None or (isinstance(raw_name, float) and raw_name != raw_name):
            continue
        name = _nfc(raw_name)
        if not name:
            continue

        # ── Ghi chú phụ (không dùng làm số trong sinh nhận xét) ──
        # Chỉ ghi lại những trường không có ý nghĩa số học rõ ràng
        remarks = ""

        # ── Thông số đo lường hiện trường — dùng để sinh nhận xét (không dùng INPS) ──
        def _cell_val(col_key: str) -> object:
            col = cm.get(col_key)
            if col is None:
                return None
            try:
                import pandas as _pd
                v = row[col]
                if v is None or (isinstance(v, float) and _pd.isna(v)):
                    return None
                return v
            except Exception:
                return None

        def _cell_str(col_key: str) -> str | None:
            v = _cell_val(col_key)
            if v is None:
                return None
            s = str(v).strip()
            return s if s else None

        excel_params = {
            "current_char": _cell_str("current_char"),
            "u_min":        _cell_val("u_min"),
            "u_max":        _cell_val("u_max"),
            "i_max":        _cell_val("i_max"),
            "delta_u":      _cell_val("delta_u"),
            "delta_i":      _cell_val("delta_i"),
            "p":            _cell_val("p"),
            "cos_phi":      _cell_val("cos_phi"),
            "thd":          _cell_val("thd"),
            "tdd":          _cell_val("tdd"),
            "pdm":          _cell_val("pdm"),
        }

        entry = {
            "name": name,
            "stt": _excel_stt(row[cm["stt"]]),
            "kind": _norm_kind(row[kind_col]),
            "nominal_voltage": None,
            "remarks": remarks,
            "excel_params": excel_params,
        }
        for mk in _metadata_keys_for_excel_name(name):
            out[mk] = entry
    return out

def _find_first_excel(root: Path) -> Path | None:
    for p in sorted(root.rglob("*")):
        if not p.is_file():
            continue
        parts = [pp.lower() for pp in p.parts]
        if any(pp == "__macosx" for pp in parts) or p.name.startswith("._"):
            continue
        if p.suffix.lower() in {".xlsx", ".xlsm", ".xls"}:
            return p
    return None

def _find_project_root(extract_root: Path) -> Path:
    """Tìm thư mục chứa các thư mục thiết bị.

    Ưu tiên một thư mục có tên ``Project_Output``; nếu không, dùng chính
    ``extract_root``.
    """
    direct = extract_root / "Project_Output"
    if direct.is_dir():
        return direct
    for p in extract_root.rglob("Project_Output"):
        if p.is_dir():
            return p
    return extract_root

def build_word_report_from_zip(
    zip_bytes: bytes,
    output_docx: str | Path,
    *,
    mba_template: str | Path | None = None,
    device_template: str | Path | None = None,
    device4_template: str | Path | None = None,
    totalmba_template: str | Path | None = None,
) -> tuple[Path, list[str]]:
    """
    Điểm nhập chính để tạo báo cáo Word từ một file ZIP chứa dữ liệu.
    
    Giải nén ZIP, tìm file Excel metadata, quét các thư mục thiết bị và 
    render báo cáo Word tổng hợp. 
    
    Quy tắc sắp xếp: MBAs luôn được đưa lên đầu danh sách (dựa trên metadata hoặc đoán tên), 
    sau đó mới đến các thiết bị đo kiểm khác. Thứ tự trong từng nhóm tuân theo 
    cột STT trong Excel.
    
    Args:
        zip_bytes: Dữ liệu nhị phân của file ZIP.
        output_docx: Đường dẫn file Word đầu ra.
        mba_template: Template cho MBA (tuỳ chọn).
        device_template: Template cho thiết bị thường (tuỳ chọn).
        device4_template: Template cho thiết bị chương 4 (tuỳ chọn).
        totalmba_template: Template cho bảng tổng hợp MBA (tuỳ chọn).

    Returns:
        tuple[Path, list[str]]: (Đường dẫn file báo cáo, danh sách các cảnh báo).
    """
    mba_template = Path(mba_template or DEFAULT_MBA_TEMPLATE)
    device_template = Path(device_template or DEFAULT_DEVICE_TEMPLATE)
    device4_template = Path(device4_template or DEFAULT_DEVICE4_TEMPLATE)
    if not mba_template.is_file():
        raise FileNotFoundError(f"Thiếu template Word MBA: {mba_template}")
    if not device_template.is_file():
        raise FileNotFoundError(f"Thiếu template Word device: {device_template}")
    if not device4_template.is_file():
        raise FileNotFoundError(f"Thiếu template Word device4: {device4_template}")

    with TemporaryDirectory(prefix="word_report_") as td:
        extract = Path(td) / "in"
        extract.mkdir()
        bio = io.BytesIO(zip_bytes)
        try:
            try:
                zf = zipfile.ZipFile(bio, "r", metadata_encoding="utf-8")
            except TypeError:
                bio.seek(0)
                zf = zipfile.ZipFile(bio, "r")
            with zf:
                zf.extractall(extract)
        except zipfile.BadZipFile as e:
            raise ValueError(f"File ZIP không hợp lệ: {e}") from e

        project_root = _find_project_root(extract)
        raw_dirs = [
            d for d in project_root.iterdir()
            if d.is_dir() and not d.name.startswith(".") and d.name != "__MACOSX"
        ]
        if not raw_dirs:
            raise ValueError(
                "ZIP không chứa thư mục thiết bị nào. Cấu trúc mong đợi: "
                "Project_Output/<Tên thiết bị>/ (a.png, PS-SDxxxx.BMP)."
            )

        excel_path = _find_first_excel(extract)
        metadata = read_device_metadata_from_excel(excel_path) if excel_path else {}

        _stt_fallback = 10**9

        def _device_dir_sort_key(p: Path) -> tuple[int, int, str]:
            if not metadata:
                kind_no_meta = _resolve_word_section_kind({}, name=p.name, default_kind=None)
                return (0 if kind_no_meta == "mba" else 1, _stt_fallback, p.name.lower())
            m = _lookup_device_metadata(metadata, p.name)
            display = _nfc(m.get("name") or p.name)
            kind = _resolve_word_section_kind({"kind": m.get("kind")}, name=display, default_kind=None)
            st = m.get("stt")
            st_val = st if isinstance(st, int) else _stt_fallback
            return (0 if kind == "mba" else 1, st_val, p.name.lower())

        device_dirs = sorted(raw_dirs, key=_device_dir_sort_key)

        devices: list[dict] = []
        for d in device_dirs:
            meta = _lookup_device_metadata(metadata, d.name)
            display = _nfc(meta.get("name") or d.name)
            devices.append({
                "name": display,
                "folder": d,
                "kind": meta.get("kind"),
                "remarks": meta.get("remarks", ""),
                "nominal_voltage": meta.get("nominal_voltage"),
                "excel_params": meta.get("excel_params") or {},
            })

        out = Path(output_docx)
        out.parent.mkdir(parents=True, exist_ok=True)
        path, warnings = build_field_word_report(
            project_root,
            out,
            mba_template=mba_template,
            device_template=device_template,
            device4_template=device4_template,
            totalmba_template=totalmba_template,
            devices=devices,
        )
        if excel_path:
            warnings.insert(0, f"Đã dùng metadata Excel: {excel_path.name}.")
        return path, warnings


def _build_chapter_from_zip(
    zip_bytes: bytes,
    output_docx: str | Path,
    *,
    chapter_filter: Literal["all", "chapter4", "chapter5"],
    device_template: str | Path | None,
    mba_template: str | Path | None = None,
    device4_template: str | Path | None = None,
    totalmba_template: str | Path | None = None,
) -> tuple[Path, list[str]]:
    """Logic dùng chung để dựng báo cáo Word (Chương 4 hoặc Chương 5) từ dữ liệu ZIP.

    Args:
        zip_bytes: Dữ liệu nhị phân file ZIP.
        output_docx: Đường dẫn file Word đầu ra.
        chapter_filter: Lọc chương cần tạo (chapter4/chapter5/all).
        device_template: Đường dẫn template thiết bị.
        mba_template: Đường dẫn template MBA.
        device4_template: Đường dẫn template thiết bị loại 4.
        totalmba_template: Đường dẫn template bảng tổng hợp MBA.

    Returns:
        tuple[Path, list[str]]: (Đường dẫn file kết quả, danh sách cảnh báo).
    """
    _mba_tpl = Path(mba_template or DEFAULT_MBA_TEMPLATE)
    _dev_tpl = Path(device_template or DEFAULT_DEVICE_TEMPLATE)
    _dev4_tpl = Path(device4_template or DEFAULT_DEVICE4_TEMPLATE)

    if chapter_filter != "chapter4" and not _mba_tpl.is_file():
        raise FileNotFoundError(f"Thiếu template Word MBA: {_mba_tpl}")
    if chapter_filter != "chapter4" and not _dev_tpl.is_file():
        raise FileNotFoundError(f"Thiếu template Word device: {_dev_tpl}")
    if chapter_filter != "chapter5" and not _dev4_tpl.is_file():
        raise FileNotFoundError(f"Thiếu template Word device4: {_dev4_tpl}")

    with TemporaryDirectory(prefix="word_chap_") as td:
        extract = Path(td) / "in"
        extract.mkdir()
        bio = io.BytesIO(zip_bytes)
        try:
            try:
                zf = zipfile.ZipFile(bio, "r", metadata_encoding="utf-8")
            except TypeError:
                bio.seek(0)
                zf = zipfile.ZipFile(bio, "r")
            with zf:
                zf.extractall(extract)
        except zipfile.BadZipFile as e:
            raise ValueError(f"File ZIP không hợp lệ: {e}") from e

        project_root = _find_project_root(extract)
        raw_dirs = [
            d for d in project_root.iterdir()
            if d.is_dir() and not d.name.startswith(".") and d.name != "__MACOSX"
        ]
        if not raw_dirs:
            raise ValueError(
                "ZIP không chứa thư mục thiết bị nào. Cấu trúc mong đợi: "
                "Project_Output/<Tên thiết bị>/ (a.png, PS-SDxxxx.BMP)."
            )

        excel_path = _find_first_excel(extract)
        metadata = read_device_metadata_from_excel(excel_path) if excel_path else {}

        _stt_fallback = 10**9

        def _device_dir_sort_key(p: Path) -> tuple[int, int, str]:
            if not metadata:
                kind_no_meta = _resolve_word_section_kind({}, name=p.name, default_kind=None)
                return (0 if kind_no_meta == "mba" else 1, _stt_fallback, p.name.lower())
            m = _lookup_device_metadata(metadata, p.name)
            display = _nfc(m.get("name") or p.name)
            kind = _resolve_word_section_kind({"kind": m.get("kind")}, name=display, default_kind=None)
            st = m.get("stt")
            st_val = st if isinstance(st, int) else _stt_fallback
            return (0 if kind == "mba" else 1, st_val, p.name.lower())

        device_dirs = sorted(raw_dirs, key=_device_dir_sort_key)

        devices: list[dict] = []
        for d in device_dirs:
            meta = _lookup_device_metadata(metadata, d.name)
            display = _nfc(meta.get("name") or d.name)
            devices.append({
                "name": display,
                "folder": d,
                "kind": meta.get("kind"),
                "remarks": meta.get("remarks", ""),
                "nominal_voltage": meta.get("nominal_voltage"),
                "excel_params": meta.get("excel_params") or {},
            })

        out = Path(output_docx)
        out.parent.mkdir(parents=True, exist_ok=True)
        path, warnings = build_field_word_report(
            project_root,
            out,
            mba_template=_mba_tpl,
            device_template=_dev_tpl,
            device4_template=_dev4_tpl,
            totalmba_template=totalmba_template,
            devices=devices,
            chapter_filter=chapter_filter,
        )
        if excel_path:
            warnings.insert(0, f"Đã dùng metadata Excel: {excel_path.name}.")
        return path, warnings


def build_chapter4_from_zip(
    zip_bytes: bytes,
    output_docx: str | Path,
    *,
    device4_template: str | Path | None = None,
) -> tuple[Path, list[str]]:
    """Tạo Chương 4 Word: chỉ các thiết bị có ``type=\"4\"`` (dùng ``device4.docx``).

    Args:
        zip_bytes: Dữ liệu ZIP đã tổ chức (output của bước \"Xử lý file sơ bộ\").
        output_docx: Đường dẫn file Word đầu ra.
        device4_template: Template cho thiết bị type-4 (tuỳ chọn, mặc định ``device4.docx``).

    Returns:
        tuple: (Đường dẫn file báo cáo, danh sách cảnh báo).
    """
    return _build_chapter_from_zip(
        zip_bytes,
        output_docx,
        chapter_filter="chapter4",
        device_template=DEFAULT_DEVICE_TEMPLATE,  # không dùng nhưng cần pass
        device4_template=device4_template,
    )


def build_chapter5_from_zip(
    zip_bytes: bytes,
    output_docx: str | Path,
    *,
    mba_template: str | Path | None = None,
    device_template: str | Path | None = None,
    totalmba_template: str | Path | None = None,
) -> tuple[Path, list[str]]:
    """Tạo Chương 5 Word: MBA + các thiết bị KHÔNG có ``type=\"4\"``.

    Giống logic cũ \"Tạo Chương 4+5\" nhưng loại bỏ các thiết bị type-4.

    Args:
        zip_bytes: Dữ liệu ZIP đã tổ chức.
        output_docx: Đường dẫn file Word đầu ra.
        mba_template, device_template, totalmba_template: Template tuỳ chọn.

    Returns:
        tuple: (Đường dẫn file báo cáo, danh sách cảnh báo).
    """
    return _build_chapter_from_zip(
        zip_bytes,
        output_docx,
        chapter_filter="chapter5",
        mba_template=mba_template,
        device_template=device_template,
        totalmba_template=totalmba_template,
    )


# ════════════════════════════════════════════════════════════════════
#   Bảng tổng hợp kết quả đo kiểm (Table 6 — table6.docx)
# ════════════════════════════════════════════════════════════════════

# Ngưỡng đánh giá cho bảng tổng hợp (theo quy tắc trong temp-ex.py)
_T6_COSPHI_LOW = 0.75
_T6_COSPHI_MID = 0.85
_T6_DELTA_I_HIGH = 10.0
_T6_TDD_HIGH = 12.0


def _t6_auto_nhan_xet(delta_I: object, cos_phi: object, tdd: object) -> str:
    """Tự động sinh đoạn văn nhận xét cho bảng tổng hợp Table 6 (Bảng 6.3).

    Args:
        delta_I: Giá trị độ lệch pha dòng điện.
        cos_phi: Giá trị hệ số công suất.
        tdd: Giá trị sóng hài dòng điện TDD.

    Returns:
        str: Câu nhận xét tổng hợp ("Vận hành ổn định" hoặc các lỗi vi phạm).
    """
    vi_pham: list[str] = []

    def _to_float(v: object) -> float | None:
        if v is None:
            return None
        try:
            import pandas as _pd
            if isinstance(v, float) and _pd.isna(v):
                return None
        except Exception:
            pass
        try:
            return float(str(v).strip().replace(",", "."))
        except (ValueError, TypeError):
            return None

    val_cos = _to_float(cos_phi)
    val_di = _to_float(delta_I)
    val_tdd = _to_float(tdd)

    if val_cos is not None and abs(val_cos) < _T6_COSPHI_LOW:
        vi_pham.append("Hệ số Cosφ còn thấp")
    if val_di is not None and val_di >= _T6_DELTA_I_HIGH:
        vi_pham.append("Độ lệch pha dòng điện còn cao")
    if val_tdd is not None and val_tdd >= _T6_TDD_HIGH:
        vi_pham.append("Tổng biến dạng sóng hài dòng điện còn cao")

    if not vi_pham:
        return "Thiết bị vận hành ổn định"
    return ", ".join(vi_pham)


def _t6_fmt(v: object, decimals: int = 2) -> str:
    """Định dạng giá trị số cho các ô trong bảng tổng hợp Table 6.

    Hỗ trợ xử lý NaN, None và định dạng dấu phẩy thập phân kiểu Việt Nam.

    Args:
        v: Giá trị cần định dạng.
        decimals: Số chữ số thập phân.

    Returns:
        str: Chuỗi đã định dạng (VD: "0,987") hoặc "—" nếu thiếu dữ liệu.
    """
    if v is None:
        return "—"
    try:
        import pandas as _pd
        if isinstance(v, float) and _pd.isna(v):
            return "—"
    except Exception:
        pass
    try:
        x = float(str(v).strip().replace(",", "."))
    except (ValueError, TypeError):
        return str(v).strip() if str(v).strip() else "—"
    if x != x:  # NaN
        return "—"
    return f"{x:.{decimals}f}".replace(".", ",")


def generate_table6_docx(
    devices: list[dict],
    output_path: str | Path,
    *,
    template_path: str | Path | None = None,
) -> Path:
    """Sinh báo cáo Word cho Bảng 6.3 (Tổng hợp kết quả đo kiểm).

    Args:
        devices: Danh sách các dict thông số thiết bị (tên, I, delta_I, cosphi, P, tdd).
        output_path: Đường dẫn lưu file .docx kết quả.
        template_path: Đường dẫn template Table 6 (tuỳ chọn).

    Returns:
        Path: Đường dẫn đến file Word đã sinh.
    """
    tpl_path = Path(template_path or DEFAULT_TABLE6_TEMPLATE)
    if not tpl_path.is_file():
        raise FileNotFoundError(f"Thiếu template Table6: {tpl_path}")

    doc = DocxTemplate(str(tpl_path))

    danh_sach: list[dict] = []
    for idx, thiet_bi in enumerate(devices, start=1):
        nhan_xet = str(thiet_bi.get("nhan_xet") or "").strip()
        if not nhan_xet:
            nhan_xet = _t6_auto_nhan_xet(
                thiet_bi.get("delta_I"),
                thiet_bi.get("cos_phi"),
                thiet_bi.get("tdd"),
            )
        danh_sach.append({
            "tt":       idx,
            "ten":      str(thiet_bi.get("ten") or "").strip(),
            "I":        _t6_fmt(thiet_bi.get("I"), 0),
            "delta_I":  _t6_fmt(thiet_bi.get("delta_I"), 1),
            "cos_phi":  _t6_fmt(thiet_bi.get("cos_phi"), 3),
            "P":        _t6_fmt(thiet_bi.get("P"), 0),
            "tdd":      _t6_fmt(thiet_bi.get("tdd"), 2),
            "nhan_xet": nhan_xet,
        })

    doc.render({"ds_thiet_bi": danh_sach}, autoescape=True)
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def generate_table6_from_zip(
    zip_bytes: bytes,
    output_path: str | Path,
    *,
    template_path: str | Path | None = None,
) -> tuple[Path, list[str]]:
    """Tạo bảng tổng hợp Table 6 (Bảng 6.3) từ file ZIP dữ liệu.

    Args:
        zip_bytes: Nội dung nhị phân file ZIP.
        output_path: Đường dẫn file Word đầu ra.
        template_path: Template Table 6 (tuỳ chọn).

    Returns:
        tuple[Path, list[str]]: (Đường dẫn file Word kết quả, danh sách các cảnh báo).
    """
    warnings: list[str] = []

    with TemporaryDirectory(prefix="table6_") as td:
        extract = Path(td) / "in"
        extract.mkdir()
        bio = io.BytesIO(zip_bytes)
        try:
            try:
                zf = zipfile.ZipFile(bio, "r", metadata_encoding="utf-8")
            except TypeError:
                bio.seek(0)
                zf = zipfile.ZipFile(bio, "r")
            with zf:
                zf.extractall(extract)
        except zipfile.BadZipFile as e:
            raise ValueError(f"File ZIP không hợp lệ: {e}") from e

        excel_path = _find_first_excel(extract)
        devices: list[dict] = []

        if excel_path is None:
            warnings.append(
                "Không tìm thấy file Excel trong ZIP — bảng tổng hợp sẽ được tạo rỗng."
            )
        else:
            try:
                import pandas as _pd
                from modules.kew.organize_field_zip import resolve_field_excel_column_map

                df = _pd.read_excel(str(excel_path), header=0, engine="openpyxl")
                if df.empty:
                    warnings.append("File Excel không có dữ liệu.")
                else:
                    try:
                        cm = resolve_field_excel_column_map(df)
                    except ValueError as e:
                        cm = None
                        warnings.append(f"Không đọc được cấu trúc Excel: {e}")

                    if cm is not None:
                        rows: list[tuple[int, dict]] = []
                        for _, row in df.iterrows():
                            raw_name = row[cm["name"]]
                            if raw_name is None or (
                                isinstance(raw_name, float) and _pd.isna(raw_name)
                            ):
                                continue
                            name = _nfc(str(raw_name).strip())
                            if not name:
                                continue

                            def _cv(col_key: str) -> object:
                                col = cm.get(col_key)
                                if col is None:
                                    return None
                                try:
                                    v = row[col]
                                    if v is None or (
                                        isinstance(v, float) and _pd.isna(v)
                                    ):
                                        return None
                                    return v
                                except Exception:
                                    return None

                            stt_raw = _excel_stt(row[cm["stt"]])
                            stt_val = stt_raw if isinstance(stt_raw, int) else 10 ** 9
                            rows.append((stt_val, {
                                "ten":     name,
                                "I":       _cv("i_max"),
                                "delta_I": _cv("delta_i"),
                                "cos_phi": _cv("cos_phi"),
                                "P":       _cv("p"),
                                "tdd":     _cv("tdd"),
                            }))

                        rows.sort(key=lambda x: x[0])
                        devices = [r[1] for r in rows]
                        warnings.insert(0, f"Đã đọc metadata Excel: {excel_path.name}.")
            except ImportError:
                warnings.append("Thiếu pandas/openpyxl — không đọc được Excel.")
            except Exception as e:
                warnings.append(f"Lỗi đọc Excel: {e}")

        out = generate_table6_docx(
            devices,
            output_path,
            template_path=template_path,
        )
        return out, warnings
